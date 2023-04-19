import docx
from docx.shared import RGBColor

# Открываем документ
doc = docx.Document('example.docx')
# Получаем первый параграф

#paragraph = doc.paragraphs[0]

p='хай'
# Создаем новый параграф для записи измененного текста
new_paragraph = doc.add_paragraph()
i=0
# Проходим по всем символам в тексте
for paragraph in doc.paragraphs:
    for char in paragraph.text:    
        if i < len(p): 
            if char==p[i]:
    # Присваиваем каждому символу отдельный цвет
                color = RGBColor(255, 0, 0)  # красный цвет
    # Создаем новый объект для символа с заданным цветом
                run = new_paragraph.add_run(char)
                run.font.color.rgb = color
                i+=1
            else:
                run = new_paragraph.add_run(char)
                i+=1
    
doc.save('example1.docx')