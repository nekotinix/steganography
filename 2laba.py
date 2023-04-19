from xml.dom.minidom import Document
from docx import *
from bitstring import BitArray
from docx.shared import RGBColor

if __name__ == "__main__":

    doc = Document('variant02.docx')
    doc1 = Document()
    bits = ''
    iteration = 0
    str_code = 'Без труда не выловишь и рыбку из пруда!'
    print('Пословица:',str_code[0:38],'\n')
    to_bytes = str_code.encode('koi8-r')
    to_bits = BitArray(to_bytes)
    print('Двоичная последовательность:',to_bits.bin,'\n')
    bits_str = to_bits.bin

    for i_par, par in enumerate(doc.paragraphs):
        print('Текст абзаца:',doc.paragraphs[i_par].text)
        new_par = doc1.add_paragraph('')
        for j_char, char in enumerate(doc.paragraphs[i_par].text):
            if iteration < len(bits_str) - 1:
                if bits_str[iteration] == '1':
                    new_par.add_run(doc.paragraphs[i_par].text[j_char]).font.color.rgb = RGBColor(0, 1, 0)
                else:
                    new_par.add_run(doc.paragraphs[i_par].text[j_char])
            else:
                new_par.add_run(doc.paragraphs[i_par].text[j_char])
            iteration += 1

    doc1.save("example1.docx")
