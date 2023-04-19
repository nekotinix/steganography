import docx
import random
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
doc = docx.Document("variant02.docx")
paragraphs = doc.paragraphs
new_paragraph = doc.add_paragraph()
message='пушка'
print('ОТ:',message)
conv_str=' '.join(format(x, 'b') for x in bytearray(message, 'koi8_r'))
print('двоичный вид сообщения koi 8:',conv_str)
conv_str=conv_str.replace(" ", "")
print('двоичный вид сообщения koi 8 без пробелов:',conv_str)
dec_mess=bytes.fromhex(hex(int(conv_str, 2))[2:]).decode(encoding="koi8_r")  
print('Проверка исходного текста кодировки:',dec_mess)
text = []
text2=[]
poz_txt=[]
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
text_v_stroke='\n'.join(text)
print(text_v_stroke)

doc = docx.Document('example1.docx')
text2=[]
for paragraph in doc.paragraphs:
    text2.append(paragraph.text)
text_v_stroke2='\n'.join(text2)
print(text_v_stroke2)



for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.text == text_v_stroke2:
            start_pos = run.start
            end_pos = run.end
            new_text = text_v_stroke
            paragraph.text = paragraph.text[:start_pos] + new_text + paragraph.text[end_pos:]



doc.save('example2.docx')

