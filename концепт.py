import docx
import random
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
doc = docx.Document("variant02.docx")
paragraphs = doc.paragraphs
new_paragraph = doc.add_paragraph()
message='пушка'
print('ОТ:',message)
conv_str=' '.join(format(x, 'b') for x in bytearray(message, 'koi8_r'))
print('двоичный вид сообщения koi 8:',conv_str)
conv_str=conv_str.replace(" ", "")
print('двоичный вид сообщения koi 8 без пробелов:',conv_str)
dec_mess=bytes.fromhex(hex(int(conv_str, 2))[2:]).decode(encoding="koi8_r")  
print('Проверка исходного текста кодировки:',dec_mess)
text = []
text2=[]
poz_txt=[]
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
text_v_stroke='\n'.join(text)
print(text_v_stroke)

new=(''.join(text2))
color2 = {}

# Создаем список ключей и значений
keys = ['1']
values = [RGBColor(255, 1, 1)]

# Заполняем словарь значениями из списков
for i in range(len(keys)):
    color2[keys[i]] = values[i]
new_character='1'
for x in range(len(new)):
    s=new[x]
    if s=='1':
        text_v_stroke=text_v_stroke.replace(text_v_stroke[x],'1',1)
        #temp=list(text_v_stroke)
        #temp[x]='1'
        #text_v_stroke="".join(temp)
        #print(x)
        #text_v_stroke=text_v_stroke[:x]+new_character+text_v_stroke[x+1:]
print('tex v str',text_v_stroke)
# Выводим результат
print('slovar2:',color2)
for char in text_v_stroke:
    if char in color2:
        # Добавляем символ с заданным цветом
        run = new_paragraph.add_run(char)
        #new_paragraph.add_run(char).font.color.rgb = colors[char]
        font = run.font
        font.color.rgb = color2[char]
        font = run.font
        font.name = 'Segoe Print'
        font.size = Pt(12)
    else:
        # Добавляем символ без изменения цвета
         run=new_paragraph.add_run(char)
         font = run.font
         font.name = 'Segoe Print'
         font.size = Pt(12)
# Сохраняем изменения в документе
doc.save('example1.docx')

        #text_v_stroke=text_v_stroke.replace(text_v_stroke[x],'1')
        #print(text_v_stroke)
    