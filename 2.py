import docx
import random
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
#colors = {
    
#    'а':RGBColor(255, 0, 0), # красный
#    'б':RGBColor(255, 0, 0), # красный
#    'в':RGBColor(255, 0, 0), # красный
#    'г':RGBColor(255, 0, 0), # красный
#    'д':RGBColor(255, 0, 0), # красный
#    'е':RGBColor(255, 0, 0), # красный
#    'ё':RGBColor(255, 0, 0), # красный
#    'ж':RGBColor(255, 0, 0), # красный
#    'з':RGBColor(255, 0, 0), # красный
#    'и':RGBColor(255, 0, 0), # красный
#    'й':RGBColor(255, 0, 0), # красный
#    'к':RGBColor(255, 0, 0), # красный
#    'л':RGBColor(255, 0, 0), # красный
#    'м':RGBColor(255, 0, 0), # красный
#    'н':RGBColor(255, 0, 0), # красный
#    'о':RGBColor(255, 0, 0), # красный
#    'п':RGBColor(255, 0, 0), # красный
#    'р':RGBColor(255, 0, 0), # красный
#    'с':RGBColor(255, 0, 0), # красный
#    'т':RGBColor(255, 0, 0), # красный
#    'у':RGBColor(255, 0, 0), # красный
#    'ф':RGBColor(255, 0, 0), # красный
#    'х':RGBColor(255, 0, 0), # красный
#    'ц':RGBColor(255, 0, 0), # красный
#    'ч':RGBColor(255, 0, 0), # красный
#    'ш':RGBColor(255, 0, 0), # красный
#    'щ':RGBColor(255, 0, 0), # красный
#    'ъ':RGBColor(255, 0, 0), # красный
#    'ы':RGBColor(255, 0, 0), # красный
#    'ь':RGBColor(255, 0, 0), # красный
#    'э':RGBColor(255, 0, 0), # красный
#    'ю':RGBColor(255, 0, 0), # красный
#    'я':RGBColor(255, 0, 0), # красный

#    'А':RGBColor(255, 0, 0), # красный
#    'Б':RGBColor(255, 0, 0), # красный
#    'В':RGBColor(255, 0, 0), # красный
#    'Г':RGBColor(255, 0, 0), # красный
#    'Д':RGBColor(255, 0, 0), # красный
#    'Е':RGBColor(255, 0, 0), # красный
#    'Ё':RGBColor(255, 0, 0), # красный
#    'Ж':RGBColor(255, 0, 0), # красный
#    'З':RGBColor(255, 0, 0), # красный
#    'И':RGBColor(255, 0, 0), # красный
#    'Й':RGBColor(255, 0, 0), # красный
#    'К':RGBColor(255, 0, 0), # красный
#    'Л':RGBColor(255, 0, 0), # красный
#    'М':RGBColor(255, 0, 0), # красный
#    'Н':RGBColor(255, 0, 0), # красный
#    'О':RGBColor(255, 0, 0), # красный
#    'П':RGBColor(255, 0, 0), # красный
#    'Р':RGBColor(255, 0, 0), # красный
#    'С':RGBColor(255, 0, 0), # красный
#    'Т':RGBColor(255, 0, 0), # красный
#    'У':RGBColor(255, 0, 0), # красный
#    'Ф':RGBColor(255, 0, 0), # красный
#    'Х':RGBColor(255, 0, 0), # красный
#    'Ц':RGBColor(255, 0, 0), # красный
#    'Ч':RGBColor(255, 0, 0), # красный
#    'Ш':RGBColor(255, 0, 0), # красный
#    'Щ':RGBColor(255, 0, 0), # красный
#    'Ъ':RGBColor(255, 0, 0), # красный
#    'Ы':RGBColor(255, 0, 0), # красный
#    'Ь':RGBColor(255, 0, 0), # красный
#    'Э':RGBColor(255, 0, 0), # красный
#    'Ю':RGBColor(255, 0, 0), # красный
#    'Я':RGBColor(255, 0, 0), # красный
#    ' ':RGBColor(255, 0, 0), # красный
#    '.':RGBColor(255, 0, 0), # красный
#    ',':RGBColor(255, 0, 0), # красный
#    '-':RGBColor(255, 0, 0), # красный
#    '\n':RGBColor(255, 0, 0), # красный
#    '!':RGBColor(255, 0, 0), # красный
#    ':':RGBColor(255, 0, 0), # красный
#    '?':RGBColor(255, 0, 0), # красный
   
#}

doc = docx.Document("variant02.docx")
paragraphs = doc.paragraphs
new_paragraph = doc.add_paragraph()
message='пушка'
print('ОТ:',message)
conv_str=' '.join(format(x, 'b') for x in bytearray(message, 'koi8_r'))
print('двоичный вид сообщения koi 8:',conv_str)
conv_str=conv_str.replace(" ", "")
print('двоичный вид сообщения koi 8 без пробелов:',conv_str)
dec_mess=bytes.fromhex(hex(int(conv_str, 2))[2:]).decode(encoding="koi8_r")  
print('Проверка исходного текста кодировки:',dec_mess)
text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
text_v_stroke='\n'.join(text)
print('text:',text)
dlina_spiska=len(text)
print('dlina teksta:',dlina_spiska-1)
#text_v_dvoichnom=' '.join(format(x, 'b') for x in bytearray(text_v_stroke, 'utf8'))
#print(text_v_dvoichnom)
text_stroka='\n'.join(text)
print('\n'.join(text))
p=[]



for i in range(len(text)):
    for j in range(len(text[i])):
        #print('dlina v elemente spiska',len(text[i]))
        print(text[i][j])
        print('i',i,'j',j)
    x=len(text[i])
    print('x=',x)
    if x!=0:
        if x== len(text[i]):
            x-=1
        number=random.randint(0,j)
        ishod=text[i][number]
        #print('nomer:',number,'i=',i)
        print('ishod=',text[i][number])
        print(i,number)
        p.append(ishod)
        
        pa=(''.join(p))
        print(pa)
# Создаем пустой словарь
colors = {}
values=[]
for o in p:
    values.append(RGBColor(255, 1, 1))
# Создаем список ключей и значений
# Заполняем словарь значениями из списков
for i in range(len(p)):
    colors[p[i]] = values[i]
# Обходим все символы в тексте пар
# Выводим результат
print(colors)



for char in text_stroka:
    if char in colors:
        # Добавляем символ с заданным цветом
        run = new_paragraph.add_run(char)
        #new_paragraph.add_run(char).font.color.rgb = colors[char]
        font = run.font
        font.color.rgb = colors[char]
        font = run.font
        font.name = 'Segoe Print'
        font.size = Pt(12)
    else:
        # Добавляем символ без изменения цвета
         run=new_paragraph.add_run(char)
         font = run.font
         font.name = 'Segoe Print'
         font.size = Pt(12)
# Сохраняем изменения в документе
doc.save('example1.docx')