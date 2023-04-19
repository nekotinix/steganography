import docx
import os
paths = []
folder = os.getcwd()
for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))
for path in paths:
    doc = docx.Document('12.docx')
i=0
colo=[]
colo2=[]
for character in doc.paragraphs:
    for run in character.runs:
        i+=1
        print('Название шрифта:', run.font.name)
        x=(str(run.font.color.rgb))
        colo.append(x)
print('Список цветов:',colo)
print('Длина списка:',len(colo))
text = []
text_pos=[]
l=0
for c  in range(len(colo)):
    if colo[c] == 'FF0000':
        print('Строка с цветом:',c)
        for paragraph in doc.paragraphs:
        
            text.append(paragraph.text)
        print('Буква:',text[c])
        text_pos.append(text[c])
           
z=''.join(text_pos) 
z=z.lower()
print('Зашифрованное слово:',z)
#x='\n'.join(text)
#doc1 = docx.Document()
#for f in x:
    #doc1.add_paragraph(f)
    #doc1.save('test.docx')
    #print(f)
import docx
from docx.shared import RGBColor

# Открываем документ
doc = docx.Document('example.docx')

# Цвет, который мы будем присваивать букве на определенной позиции
color = RGBColor(255, 0, 0)  # Красный цвет

# Перебираем все абзацы в документе
for paragraph in doc.paragraphs:
    # Получаем текст абзаца
    text = paragraph.text
    
    # Определяем позицию буквы, которую мы будем выделять цветом
    position = 1  # Здесь мы выбрали пятую букву в каждом абзаце
    
    # Если абзац содержит нужную позицию, то выделяем букву цветом
    if len(text) > position:
        run = paragraph.runs[position]
        font = run.font
        font.color.rgb = color