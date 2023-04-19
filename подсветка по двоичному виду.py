import docx
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt


doc = docx.Document("variant02.docx")
paragraphs = doc.paragraphs
new_paragraph = doc.add_paragraph()
message='пушка'
print('ОТ:',message)
conv_str=' '.join(format(x, 'b') for x in bytearray(message, 'koi8_r'))
print('двоичный вид сообщения koi 8:',conv_str)
conv_str=conv_str.replace(" ", "")
print('двоичный вид сообщения koi 8 без пробелов:',conv_str)
dec_mess=bytes.fromhex(hex(int(conv_str, 2))[2:]).decode(encoding="koi8_r")  
print('Проверка исходного текста кодировки:',dec_mess)
text = []
text2=[]
poz_txt=[]
matrica=[]

for paragraph in doc.paragraphs:
    text.append(paragraph.text)
text_v_stroke='\n'.join(text)

for i in range(len(text_v_stroke)):
    conv_str2=text_v_stroke[i].replace(text_v_stroke[i], "0")
    text2.append(conv_str2)
new_str2=''.join(text2)
print('pustaya stroka',new_str2)
print('Матрица текста',text2)
for word in text:
    new_word = ''
    for letter in word:
        new_letter = '0' # заменяем букву на её верхний регистр
        new_word += new_letter # добавляем измененную букву в новое слово
    matrica.append(new_word)
print('matrica',matrica)


new_str=''
const=0
vosmerka=[]
for i in poz_txt:
    new_str=conv_str[const:i]
    print('poza',poz_txt,new_str)
    for z in new_str:
        vosmerka.append(z)
print('spisok s kodom',vosmerka)
new_str3=''.join(vosmerka)
print('novoe',new_str3)
for i in range(len(vosmerka)):
    text2[i]=vosmerka[i]
print('t2',text2)

print('text:',text)
dlina_spiska=len(text)
print('dlina texta:',text_v_stroke)
print('dlina teksta:',dlina_spiska)
text_stroka='\n'.join(text)
print('text:','\n'.join(text))
p=[]
list_dlina=[]    

listj=[]
listnew=[]
for i in range(len(text)):
    for j in range(len(text[i])):
        print('dlina v elemente spiska',len(text[i]))
        print('text',text[i][j],text2[i])
        zzs=len(text[i])+1
        listnew.append(zzs)
        dlina=list(dict.fromkeys(listnew))
        print('i',i,'j',j)
        
        x=len(text[i])
        print('x=',x)
        listj.append(j)
        print('listj:',listj)
    if x!=0:
         if x== len(text[i]):
             x-=1
             #if text2[i]=='1':
             print('t222',text2[i])  
             number=j
             print('nomerj',number,'i',i,'j',j)
             ishod=text[i][number]
                 #listj.append(number)
                    
                 #print('listj:',listj)
             print('nomer:',number,'i=',i)
             print('ishod=',text[i][number])
             p.append(ishod)
             conv_str2=text_v_stroke.replace(text[i][number], "1",1)
             text_v_stroke=conv_str2
    listj.clear()
        
       
    pa=(''.join(p))
    print('pa',pa)
print('stroka',text_v_stroke)
const4=-1
text4=[]
string=''.join(text2)
for u in dlina:
    for v in range(u):
        if v==0:
            r=0
            ds=string[r:u]
            const4+=1
            text4.append(ds)
print('text4',text4)
search_item = ''
positions = []



for i, item in enumerate(text):
    if item == search_item:
        positions.append(i)

print('posic:',positions) 
for i in range(len(text)):
    for j in range(len(text[i])):
        
        listnew.append(len(text[i]))
        
        print('i',i,'j',j)
        if text4[i][j]=='1':
            print('4',text)
            text_stroka.replace(text_stroka,'1')
values=[]
for o in p:
    print(o)
    values.append(RGBColor(255, 1, 1))
# Создаем список ключей и значений
# Заполняем словарь значениями из списков

color2={}
# Создаем список ключей и значений
keys = ['1']
values = [RGBColor(255, 1, 1)]

# Заполняем словарь значениями из списков
for i in range(len(keys)):
    color2[keys[i]] = values[i]

# Выводим результат
print('slovar2:',color2)


for char in text_v_stroke:
    if char in color2:
        # Добавляем символ с заданным цветом
        run = new_paragraph.add_run(char)
        #new_paragraph.add_run(char).font.color.rgb = colors[char]
        font = run.font
        font.color.rgb = color2[char]
        font = run.font
        font.name = 'Segoe Print'
        font.size = Pt(12)
    else:
        # Добавляем символ без изменения цвета
         run=new_paragraph.add_run(char)
         font = run.font
         font.name = 'Segoe Print'
         font.size = Pt(12)
# Сохраняем изменения в документе
doc.save('example1.docx')
#print('tvs',text_v_stroke3,len(text_v_stroke))





#for paragraph in doc.paragraphs:
#    for run in paragraph.runs:
#        if run.text == text_v_stroke:
#            start_pos = run.start
#            end_pos = run.end
#            new_text = text_v_stroke3
#            paragraph.text = paragraph.text[:start_pos] + new_text + paragraph.text[end_pos:]
#    doc.save('example2.docx')